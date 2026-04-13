"""
services.py — Lógica de procesamiento PDF
Librerías: PyPDF2, PyMuPDF (fitz), pdf2docx, python-docx, openpyxl, reportlab
"""

import io
import os
import zipfile

import fitz  # PyMuPDF
import openpyxl
import PyPDF2
from docx import Document
from openpyxl.styles import Font, PatternFill
from pdf2docx import Converter
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Image as RLImage
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

try:
    import pdfplumber
except Exception:  # pragma: no cover - dependencia opcional
    pdfplumber = None

# ── Utilidades ────────────────────────────────────────────────────

def contar_paginas(ruta_pdf):
    with open(ruta_pdf, 'rb') as f:
        reader = _crear_reader(f)
        return len(reader.pages)


def _crear_reader(file_obj, contrasena=''):
    """Abre un PDF y, si esta cifrado, intenta descifrarlo."""
    reader = PyPDF2.PdfReader(file_obj)
    if reader.is_encrypted:
        try:
            # Si llega contraseña, se intenta primero con esa.
            resultado = reader.decrypt(contrasena or '')
            if resultado == 0 and contrasena:
                # Fallback para PDFs con contraseña vacía.
                resultado = reader.decrypt('')
        except Exception as exc:
            raise ValueError('El archivo PDF esta cifrado y requiere contraseña para procesarse') from exc
        if resultado == 0:
            raise ValueError('El archivo PDF esta cifrado y requiere contraseña para procesarse')
    return reader


# ── Unión ─────────────────────────────────────────────────────────

def unir_pdfs(rutas, contrasena=''):
    writer = PyPDF2.PdfWriter()
    for ruta in rutas:
        with open(ruta, 'rb') as f:
            reader = _crear_reader(f, contrasena=contrasena)
            for page in reader.pages:
                writer.add_page(page)
    buf = io.BytesIO()
    writer.write(buf)
    return buf.getvalue()


# ── División ──────────────────────────────────────────────────────

def dividir_pdf(ruta, inicio, fin):
    with open(ruta, 'rb') as f:
        reader = _crear_reader(f)
        total = len(reader.pages)
        inicio = max(1, inicio)
        fin = min(fin, total)
        if inicio > fin:
            raise ValueError(f'Rango inválido: {inicio}-{fin} (total {total} páginas)')
        writer = PyPDF2.PdfWriter()
        for i in range(inicio - 1, fin):
            writer.add_page(reader.pages[i])
    buf = io.BytesIO()
    writer.write(buf)
    return buf.getvalue()


def dividir_pdf_por_modo(ruta, modo):
    with open(ruta, 'rb') as f:
        reader = _crear_reader(f)
        total = len(reader.pages)
        if total == 0:
            raise ValueError('El PDF no contiene paginas')

        if modo == 'todas':
            out = io.BytesIO()
            with zipfile.ZipFile(out, 'w', compression=zipfile.ZIP_DEFLATED) as zf:
                for i in range(total):
                    writer = PyPDF2.PdfWriter()
                    writer.add_page(reader.pages[i])
                    page_buf = io.BytesIO()
                    writer.write(page_buf)
                    zf.writestr(f'pagina_{i + 1}.pdf', page_buf.getvalue())
            return out.getvalue(), 'zip'

        if modo == 'pares':
            indices = [i for i in range(total) if (i + 1) % 2 == 0]
        elif modo == 'impares':
            indices = [i for i in range(total) if (i + 1) % 2 == 1]
        else:
            raise ValueError(f'Modo de division no soportado: {modo}')

        if not indices:
            raise ValueError('No hay paginas para el modo seleccionado')

        writer = PyPDF2.PdfWriter()
        for i in indices:
            writer.add_page(reader.pages[i])
        buf = io.BytesIO()
        writer.write(buf)
        return buf.getvalue(), 'pdf'


# ── Conversión ────────────────────────────────────────────────────

def pdf_a_docx(ruta_pdf):
    ruta_out = ruta_pdf.replace('.pdf', '_convertido.docx')
    cv = Converter(ruta_pdf)
    cv.convert(ruta_out, start=0, end=None)
    cv.close()
    with open(ruta_out, 'rb') as f:
        data = f.read()
    os.remove(ruta_out)
    return data


def pdf_a_xlsx(ruta_pdf):
    doc = fitz.open(ruta_pdf)
    tables_por_pagina = {}

    # Extractor especializado (si está disponible): mejora tablas complejas.
    if pdfplumber is not None:
        try:
            settings = {
                'vertical_strategy': 'text',
                'horizontal_strategy': 'text',
                'snap_tolerance': 3,
                'join_tolerance': 3,
                'intersection_tolerance': 3,
            }
            with pdfplumber.open(ruta_pdf) as pdf_plumber:
                for page_num, page in enumerate(pdf_plumber.pages, start=1):
                    raw_tables = page.extract_tables(table_settings=settings) or []
                    cleaned_tables = []
                    for raw_table in raw_tables:
                        cleaned_rows = []
                        for row in raw_table or []:
                            vals = [str(c).strip() if c is not None else '' for c in row]
                            if any(vals):
                                cleaned_rows.append(vals)
                        if cleaned_rows:
                            cleaned_tables.append(cleaned_rows)
                    if cleaned_tables:
                        tables_por_pagina[page_num] = cleaned_tables
        except Exception:
            tables_por_pagina = {}

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = 'Contenido PDF'
    ws.column_dimensions['A'].width = 10
    ws.column_dimensions['B'].width = 8
    ws.column_dimensions['C'].width = 8

    def _tables_from_page(page):
        # 1) Intento principal: detector de tablas nativo de PyMuPDF.
        # Es más preciso para PDFs digitales con líneas/estructura tabular.
        try:
            found = page.find_tables()
            if found and found.tables:
                extracted = []
                for tbl in found.tables:
                    rows = tbl.extract() or []
                    cleaned_rows = []
                    for row in rows:
                        vals = [str(c).strip() if c is not None else '' for c in row]
                        if any(vals):
                            cleaned_rows.append(vals)
                    if cleaned_rows:
                        extracted.append(cleaned_rows)
                if extracted:
                    return extracted
        except Exception:
            pass

        # 2) Fallback heurístico por coordenadas de palabras.
        # word tuple: (x0, y0, x1, y1, text, block_no, line_no, word_no)
        words = page.get_text('words')
        if not words:
            return []

        words = sorted(words, key=lambda w: (round(w[1], 1), w[0]))
        row_tol = 3.0
        rows = []
        current = []
        current_y = None

        for w in words:
            if current_y is None or abs(w[1] - current_y) <= row_tol:
                current.append(w)
                current_y = w[1] if current_y is None else (current_y + w[1]) / 2
            else:
                rows.append(sorted(current, key=lambda x: x[0]))
                current = [w]
                current_y = w[1]
        if current:
            rows.append(sorted(current, key=lambda x: x[0]))

        parsed_rows = []
        for row_words in rows:
            cells = []
            row_y = row_words[0][1]
            cell_text = row_words[0][4]
            cell_x0 = row_words[0][0]
            prev_x1 = row_words[0][2]

            # Agrupa palabras cercanas como una celda; brechas amplias crean nueva celda.
            for w in row_words[1:]:
                gap = w[0] - prev_x1
                if gap >= 6:
                    cells.append((cell_x0, cell_text.strip()))
                    cell_text = w[4]
                    cell_x0 = w[0]
                else:
                    cell_text += ' ' + w[4]
                prev_x1 = w[2]
            cells.append((cell_x0, cell_text.strip()))
            parsed_rows.append({'y': row_y, 'cells': cells})

        # Segmenta múltiples tablas por saltos verticales grandes.
        tables_raw = []
        if parsed_rows:
            current = [parsed_rows[0]]
            for row_info in parsed_rows[1:]:
                y_gap = row_info['y'] - current[-1]['y']
                if y_gap > 28:
                    tables_raw.append(current)
                    current = [row_info]
                else:
                    current.append(row_info)
            if current:
                tables_raw.append(current)

        def build_grid(rows_group):
            # Construye columnas estables por posición X para cada tabla.
            col_centers = []
            x_tol = 10.0

            def assign_col(x):
                for idx, cx in enumerate(col_centers):
                    if abs(x - cx) <= x_tol:
                        col_centers[idx] = (cx + x) / 2
                        return idx
                col_centers.append(x)
                col_centers.sort()
                return col_centers.index(x)

            grid = []
            for row_info in rows_group:
                row_vals = {}
                for x, txt in row_info['cells']:
                    cidx = assign_col(x)
                    row_vals[cidx] = (row_vals.get(cidx, '') + ' ' + txt).strip()
                max_col = max(row_vals) if row_vals else -1
                grid.append([row_vals.get(ci, '') for ci in range(max_col + 1)])

            if not grid:
                return grid

            # Normaliza filas y elimina columnas casi vacías (ruido por desalineación).
            max_cols = max(len(r) for r in grid)
            norm = [r + [''] * (max_cols - len(r)) for r in grid]
            col_non_empty = [sum(1 for row in norm if str(row[c]).strip()) for c in range(max_cols)]

            row_count = len(norm)
            keep_cols = [c for c, cnt in enumerate(col_non_empty) if cnt > 0]

            compact = []
            for row in norm:
                compact.append([row[c] for c in keep_cols])

            # Elimina columnas vacías de borde (izquierda/derecha) que solo generan desplazamiento.
            while compact and compact[0] and all(not str(r[0]).strip() for r in compact):
                compact = [r[1:] for r in compact]
            while compact and compact[0] and all(not str(r[-1]).strip() for r in compact):
                compact = [r[:-1] for r in compact]

            return compact

        return [build_grid(group) for group in tables_raw if group]

    def _non_empty_count(row):
        return sum(1 for c in row if str(c).strip())

    def _looks_like_header(row):
        if len(row) < 2:
            return False
        non_empty = [c for c in row if str(c).strip()]
        if len(non_empty) < 2:
            return False
        alpha_like = 0
        for cell in non_empty:
            txt = str(cell)
            has_digit = any(ch.isdigit() for ch in txt)
            if not has_digit:
                alpha_like += 1
        return alpha_like >= max(2, int(len(non_empty) * 0.6))

    def _normalize_row_to_header(row, expected_cols):
        """Ajusta filas a la cantidad de columnas del encabezado."""
        vals = [str(v).strip() for v in row]

        # 1) Une fragmentos monetarios partidos por token de coma decimal/miles.
        merged = []
        i = 0
        while i < len(vals):
            cur = vals[i]
            nxt = vals[i + 1] if i + 1 < len(vals) else ''
            if nxt and nxt.startswith(',') and cur:
                merged.append((cur + nxt).replace(' ', ''))
                i += 2
                continue
            merged.append(cur)
            i += 1

        # 2) Si sobran columnas, intenta compactar hacia la izquierda.
        while len(merged) > expected_cols:
            # prioriza fusionar tokens pequeños/ruidosos.
            idx = None
            for j in range(1, len(merged)):
                t = merged[j]
                if len(t) <= 2 or t.startswith(','):
                    idx = j
                    break
            if idx is None:
                idx = len(merged) - 1
            merged[idx - 1] = (merged[idx - 1] + merged[idx]).strip()
            del merged[idx]

        if len(merged) < expected_cols:
            merged.extend([''] * (expected_cols - len(merged)))

        return merged

    all_rows = []
    detected_header = None
    max_data_cols = 1
    for i, page in enumerate(doc, start=1):
        tables = tables_por_pagina.get(i) or _tables_from_page(page)
        if not tables:
            texto = page.get_text('text').strip()
            all_rows.append((i, 1, 1, [texto]))
            max_data_cols = max(max_data_cols, 1)
            continue

        for tnum, rows in enumerate(tables, start=1):
            start_idx = 0
            header_idx = None
            expected_cols = None
            for idx, row in enumerate(rows):
                if _looks_like_header(row):
                    header_idx = idx
                    break

            if header_idx is not None:
                expected_cols = len(rows[header_idx])
                if detected_header is None:
                    detected_header = rows[header_idx]
                start_idx = header_idx + 1

            out_row_num = 1
            for row in rows[start_idx:]:
                if expected_cols:
                    row = _normalize_row_to_header(row, expected_cols)

                non_empty = _non_empty_count(row)
                # Evita meter texto narrativo suelto en tablas de datos.
                if non_empty == 0:
                    continue
                if non_empty == 1:
                    text = str(next((c for c in row if str(c).strip()), '')).strip()
                    if len(text) > 45 and not any(ch.isdigit() for ch in text):
                        continue

                all_rows.append((i, tnum, out_row_num, row))
                out_row_num += 1
                max_data_cols = max(max_data_cols, len(row) if row else 1)

    if detected_header:
        header = ['Página', 'Tabla', 'Fila', *detected_header]
        if len(detected_header) < max_data_cols:
            header.extend([f'Col{c}' for c in range(len(detected_header) + 1, max_data_cols + 1)])
    else:
        header = ['Página', 'Tabla', 'Fila', *[f'Col{c}' for c in range(1, max_data_cols + 1)]]

    ws.append(header)
    for page_num, table_num, row_num, row_vals in all_rows:
        ws.append([page_num, table_num, row_num, *row_vals])

    # Ajuste de ancho básico para columnas de datos.
    for c in range(4, 4 + max_data_cols):
        col_letter = openpyxl.utils.get_column_letter(c)
        ws.column_dimensions[col_letter].width = 26

    # Encabezado fijo y con formato para navegación rápida.
    ws.freeze_panes = 'A2'
    ws.auto_filter.ref = f"A1:{openpyxl.utils.get_column_letter(3 + max_data_cols)}{ws.max_row}"
    header_fill = PatternFill(start_color='1B4FD8', end_color='1B4FD8', fill_type='solid')
    for cell in ws[1]:
        cell.font = Font(bold=True, color='FFFFFF')
        cell.fill = header_fill

    doc.close()
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def docx_a_pdf(ruta_docx):
    doc = Document(ruta_docx)
    buf = io.BytesIO()
    pdf_doc = SimpleDocTemplate(buf, pagesize=A4,
                                leftMargin=50, rightMargin=50,
                                topMargin=50, bottomMargin=50)
    styles = getSampleStyleSheet()
    story = []

    # Mapa para recorrer el documento respetando el orden real de párrafos.
    paragraph_by_element = {para._p: para for para in doc.paragraphs}
    max_image_width = A4[0] - 100  # 50 pt de margen a cada lado

    for element in doc.element.body.iterchildren():
        if element.tag.endswith('}p'):
            para = paragraph_by_element.get(element)
            if para is None:
                continue

            texto = para.text.strip()
            if texto:
                style = styles['Heading1'] if para.style.name.startswith('Heading') else styles['Normal']
                story.append(Paragraph(texto, style))
                story.append(Spacer(1, 4))

            # Extrae imágenes inline del párrafo y las agrega al PDF.
            for run in para.runs:
                for blip in run._element.xpath('.//a:blip'):
                    rid = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                    if not rid:
                        continue
                    image_part = para.part.related_parts.get(rid)
                    if not image_part:
                        continue

                    img = RLImage(io.BytesIO(image_part.blob))
                    if img.drawWidth > max_image_width:
                        scale = max_image_width / img.drawWidth
                        img.drawWidth = max_image_width
                        img.drawHeight = img.drawHeight * scale

                    story.append(img)
                    story.append(Spacer(1, 8))

            if not texto and not para.runs:
                story.append(Spacer(1, 8))

    pdf_doc.build(story)
    return buf.getvalue()


def xlsx_a_pdf(ruta_xlsx):
    wb = openpyxl.load_workbook(ruta_xlsx)
    ws = wb.active
    buf = io.BytesIO()
    pdf_doc = SimpleDocTemplate(buf, pagesize=A4,
                                leftMargin=30, rightMargin=30,
                                topMargin=40, bottomMargin=40)
    data = []
    for row in ws.iter_rows(values_only=True):
        data.append([str(cell) if cell is not None else '' for cell in row])
    if not data:
        data = [['(hoja vacía)']]
    col_count = max(len(r) for r in data)
    col_width = max(30, int(500 / col_count))
    table = Table(data, colWidths=[col_width] * col_count, repeatRows=1)
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1B4FD8')),
        ('TEXTCOLOR',  (0, 0), (-1, 0), colors.white),
        ('FONTSIZE',   (0, 0), (-1, 0), 9),
        ('FONTSIZE',   (0, 1), (-1, -1), 8),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#EBF2FF')]),
        ('GRID',       (0, 0), (-1, -1), 0.3, colors.HexColor('#C3D9FF')),
        ('VALIGN',     (0, 0), (-1, -1), 'TOP'),
    ]))
    pdf_doc.build([table])
    return buf.getvalue()


# ── Compresión ────────────────────────────────────────────────────

def comprimir_pdf(ruta_pdf, nivel='media'):
    """Comprime un PDF preservando contenido vectorial y evitando páginas en blanco."""
    doc = fitz.open(ruta_pdf)
    try:
        opts = {
            'baja': {
                'garbage': 1,
                'deflate': True,
                'deflate_images': False,
                'deflate_fonts': False,
                'clean': False,
            },
            'media': {
                'garbage': 3,
                'deflate': True,
                'deflate_images': True,
                'deflate_fonts': True,
                'clean': True,
            },
            'alta': {
                'garbage': 4,
                'deflate': True,
                'deflate_images': True,
                'deflate_fonts': True,
                'clean': True,
            },
        }.get(nivel, {
            'garbage': 3,
            'deflate': True,
            'deflate_images': True,
            'deflate_fonts': True,
            'clean': True,
        })

        # Compresión nativa de PDF: conserva mejor texto/imágenes y evita blanco total.
        return doc.tobytes(**opts)
    finally:
        doc.close()


# ── Cifrado ───────────────────────────────────────────────────────

def cifrar_pdf(ruta_pdf, contrasena):
    with open(ruta_pdf, 'rb') as f:
        reader = _crear_reader(f)
        writer = PyPDF2.PdfWriter()
        for page in reader.pages:
            writer.add_page(page)
    writer.encrypt(
        user_password=contrasena,
        owner_password=contrasena,
        use_128bit=False,
    )
    buf = io.BytesIO()
    writer.write(buf)
    return buf.getvalue()
